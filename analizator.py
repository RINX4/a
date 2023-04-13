import string
from collections import Counter  # for counting elms and n-grams
import docx  # docx
import fitz  # pdf
import nltk  # language processing
import pandas as pd  # csv and xls
import pymorphy2  # lemmatization
from tkinter import filedialog  # interface
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from tkinter.ttk import Checkbutton
from tkinter.ttk import Combobox
from tkinter import scrolledtext
from matplotlib import pyplot as plt  # graphics
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from nltk import ngrams  # ngrams
from nltk import word_tokenize  # tokenization
from nltk.corpus import stopwords  # deleting spam
from nltk.probability import FreqDist  # distribution


def open_source_file(file_name):  # открывает исходный файл
	if file_name.endswith('txt'):
		with open(f'{file_name}', "r", encoding='UTF-8') as source_file_data:
			raw_source_text = source_file_data.read()
			return raw_source_text
	elif file_name.endswith('pdf'):
		doc = fitz.open(file_name)
		text = [page.get_text("text") for page in doc]
		raw_source_text = "\n".join(text)
		return raw_source_text
	elif file_name.endswith('docx'):
		source_file_data = docx.Document(file_name)
		text = [paragraph.text for paragraph in source_file_data.paragraphs]
		raw_source_text = "\n".join(text)
		return raw_source_text


def raw_text_processing(raw_source_text):  # избавляет текст от пунктуации
	raw_text_length = len(raw_source_text)
	raw_source_text = raw_source_text.lower()
	spec_char = string.punctuation + '\n\xa0«»\t—–-…“„'  # добавляем нужные символы
	raw_source_text = "".join([char for char in raw_source_text if char not in spec_char])  # нет пунктуации
	cleared_text = "".join([char for char in raw_source_text if char not in string.digits])  # нет чисел
	return cleared_text, raw_text_length


def tokenization(cleared_text):  # выделяет слова из очищенной "массы"
	text_tokens = word_tokenize(cleared_text)
	tokenized_text = nltk.Text(text_tokens)
	return tokenized_text, text_tokens


def spam_words_delete(text_tokens):  # удаляет стоп-слова, слова паразиты и все бесполезное
	russian_stopwords = stopwords.words("russian")
	russian_stopwords.extend(['это', 'нею', 'б', 'ах'])  # добавляем в список стопслов свои
	no_spam_tokens = []
	for token in text_tokens:
		if token not in russian_stopwords:
			no_spam_tokens.append(token)
	return no_spam_tokens


def lemmatization(no_spam_tokens):  # приводит слова к нормальной, "словарной" форме
	lemmatized_tokens = []
	single_lemmatized_tokens = []
	morph = pymorphy2.MorphAnalyzer()
	for word in no_spam_tokens:  # запускаем проверку слов
		parsed_word = morph.parse(word)[0]
		lemmatized_tokens.append(parsed_word.normal_form)
	for word in lemmatized_tokens:
		if word not in single_lemmatized_tokens:
			single_lemmatized_tokens.append(word)
	return lemmatized_tokens, single_lemmatized_tokens


def storing_to_dataframe(lemmatized_tokens, single_lemmatized_tokens,
                         saved_file_name):  # записываем результаты в таблицу
	lemmatized_data = dict(col1=lemmatized_tokens)
	single_lemmatized_data = dict(col1=single_lemmatized_tokens)
	single_lemmatized_dataframe = pd.DataFrame(single_lemmatized_data)
	single_lemmatized_dataframe.to_excel(f'{saved_file_name}\\Список слов без повторений.xlsx', index=False)
	lemmatized_dataframe = pd.DataFrame(lemmatized_data)
	lemmatized_dataframe.to_excel(f'{saved_file_name}\\Список слов.xlsx', index=False)


def get_most_common_words(text):
	freq_of_dist = FreqDist(text)
	most_common_word = (freq_of_dist.most_common(1))[0]
	return most_common_word, freq_of_dist


def frequency_plotter(freq_of_dist, saved_file_name, graph_name):
	freq_of_dist = dict(freq_of_dist)  # создаем таблицу со всеми словами и частотой употребления
	sorted_freq_of_dist = {}
	sorted_keys = sorted(freq_of_dist, key=freq_of_dist.get, reverse=True)
	for w in sorted_keys:
		sorted_freq_of_dist[w] = freq_of_dist[w]
	f = list(sorted_freq_of_dist.items())[:30]
	first30words = []
	for pair in f:
		first30words.append(pair[0])
		first30words.append(pair[1])
	x = [v for k, v in enumerate(first30words) if not k % 2]
	y = [v for k, v in enumerate(first30words) if k % 2]
	plt.xlabel("Слова")
	plt.ylabel("Количество употреблений")
	plt.grid()
	plt.plot(x, y)
	plt.xticks(rotation=90)
	plt.tight_layout()
	plt.autoscale()
	plt.savefig(f'{saved_file_name}\\{graph_name}')
	plt.clf()


def main(file_name, saved_file_name):
	def data_to_docx():
		data_out = docx.Document()
		data_out.add_paragraph("Характеристики текста:")
		data_out.add_paragraph(f'Символов до обработки: {len(raw_source_text)}')
		data_out.add_paragraph(f'Символов после удаления пунктуации: {len(cleared_text)}')
		data_out.add_paragraph(f'Знаков пунктуации: {len(raw_source_text) - len(cleared_text)}')
		data_out.add_paragraph(f'Текст содержит слов: {len(text_tokens)}')
		data_out.add_paragraph(f'Самое частое слово до уборки стоп-слов:'
		                       f' "{most_common_word[0]}" в количестве {most_common_word[1]}')
		data_out.add_paragraph(f'Самое частое слово после уборки стоп-слов:'
		                       f'"{cleared_most_common_word[0]}" в количестве {cleared_most_common_word[1]}')
		data_out.add_paragraph(f'Слов после удаления стоп-слов: {len(no_spam_tokens)}')
		data_out.add_paragraph(f'Удалено стоп-слов: {len(text_tokens) - len(no_spam_tokens)}')
		data_out.add_paragraph(
			f'При лемматизации удалено слов: {len(no_spam_tokens) - len(single_lemmatized_tokens)}')
		data_out.add_paragraph(f'Уникальных слов: {len(single_lemmatized_tokens)}')
		data_out.add_paragraph("Графики:")
		data_out.add_paragraph("График 1. Слова до уборки стоп-слов")
		frequency_plotter(freq_of_dist,saved_file_name, 'График 1. Слова до уборки стоп-слов.png')
		data_out.add_picture(f'{saved_file_name}\\График 1. Слова до уборки стоп-слов.png')
		frequency_plotter(cleared_freq_of_dist, saved_file_name, 'График 2. Слова после уборки стоп-слов.png')
		data_out.add_paragraph("График 2. Слова после уборки стоп-слов")
		data_out.add_picture(f'{saved_file_name}\\График 2. Слова после уборки стоп-слов.png')
		data_out.save(f'{saved_file_name}\\Результаты анализа.docx')
	
	raw_source_text = open_source_file(file_name)
	cleared_text = raw_text_processing(raw_source_text)[0]
	tokenized_text, text_tokens = tokenization(cleared_text)
	no_spam_tokens = spam_words_delete(text_tokens)
	lemmatized_tokens, single_lemmatized_tokens = lemmatization(no_spam_tokens)
	storing_to_dataframe(lemmatized_tokens, single_lemmatized_tokens, saved_file_name)
	most_common_word, freq_of_dist = get_most_common_words(tokenized_text)
	cleared_most_common_word, cleared_freq_of_dist = get_most_common_words(no_spam_tokens)
	return data_to_docx, raw_source_text, tokenized_text, no_spam_tokens, lemmatized_tokens


def ngrams_cal(save_file_name, text_input, ngram_count, word_or_symbol):
	if word_or_symbol == 'Анализ символов':
		united_text = ''.join(text_input)
		raw_grams = ngrams(united_text, ngram_count)
		counted_ngrams = Counter(raw_grams)
		df = pd.DataFrame(counted_ngrams.keys())
		df["Количество употреблений"] = counted_ngrams.values()
		df.to_excel(f'{save_file_name}\\Результат подсчета n-грамм.xlsx')
	elif word_or_symbol == 'Анализ слов':
		raw_grams = ngrams(text_input, ngram_count)
		counted_ngrams = Counter(raw_grams)
		df = pd.DataFrame(counted_ngrams.keys())
		df["Количество употреблений"] = counted_ngrams.values()
		df.to_excel(f'{save_file_name}\\Результат подсчета n-грамм.xlsx')


def interface():
	def start():
		def ngramm_window_show():
			def choice():
				which_ngram = analysis_type.get()
				ngram_count = int(how_many_ngram.get())
				word_or_symbol = word_or_symbol_choice.get()
				pass_ngram_param(saved_file_name, which_ngram, ngram_count, word_or_symbol)
				messagebox.showinfo('Уведомление', 'Этап подсчета n-грамм выполнен')
				messagebox.showinfo('Уведомление', 'Работа завершена \nМожно выходить')
				ngram_window.destroy()
				if res == 'yes':
					show_results(saved_file_name)
			
			def pass_ngram_param(saved_file_name, which_ngram, ngram_count, word_or_symbol):
				if which_ngram == 'Оригинальный текст':
					ngrams_cal(saved_file_name, program[1], ngram_count, word_or_symbol)
				elif which_ngram == 'Без пунктуации и цифр':
					ngrams_cal(saved_file_name, program[2], ngram_count, word_or_symbol)
				elif which_ngram == 'Без стоп-слов':
					ngrams_cal(saved_file_name, program[3], ngram_count, word_or_symbol)
				elif which_ngram == 'Лемматизованный текст':
					ngrams_cal(saved_file_name, program[4], ngram_count, word_or_symbol)
				else:
					ngram_window.destroy()
					messagebox.showinfo('Уведомление', 'Неверно указаны вводные данные')
					ngramm_window_show()
			
			ngram_window = Tk()
			ngram_window.title("Анализ n-грамм")
			ngram_label = Label(ngram_window, text='Выберите параметры анализа\n и степень n-граммы',
			                    font=("Arial Bold", 14))
			analysis_type = Combobox(ngram_window)
			analysis_type['values'] = ['Оригинальный текст', 'Без пунктуации и цифр',
			                           'Без стоп-слов', 'Лемматизованный текст']
			analysis_type.current(0)
			word_or_symbol_choice = Combobox(ngram_window)
			word_or_symbol_choice['values'] = ['Анализ символов', 'Анализ слов']
			word_or_symbol_choice.current(0)
			choice_button = Button(ngram_window, text="Запустить анализ n-грамм", command=choice, padx=5, pady=5, bd=0,
			                       fg='#fff',
			                       bg='#08f', underline=0, activebackground='#fff', activeforeground='#fff',
			                       cursor='hand2')
			how_many_ngram = Combobox(ngram_window)
			how_many_ngram['values'] = [i for i in range(2, 11)]
			how_many_ngram.current(0)
			ngram_label.grid(column=1, row=0, padx=10, pady=10)
			analysis_type.grid(column=0, row=1, padx=10, pady=10)
			word_or_symbol_choice.grid(column=1, row=1, padx=10, pady=10)
			how_many_ngram.grid(column=2, row=1, padx=10, pady=10)
			choice_button.grid(column=1, row=2, padx=10, pady=10)
			
			def btn_focus_in(e=None):
				choice_button.configure(fg='#08f')
				choice_button.configure(bg='#fff')
			
			def btn_focus_out(e=None):
				choice_button.configure(bg='#08f')
				choice_button.configure(fg='#fff')
			
			choice_button.bind('<Enter>', btn_focus_in)
			choice_button.bind('<Leave>', btn_focus_out)
			
			ngram_window.mainloop()
		
		try:
			res = messagebox.askquestion('', 'Нужно ли выводить результаты на экран?')
			file_name = "{}".format(file_name_entry.get())
			saved_file_name = "{}".format(save_directory_entry.get())
			check1 = chk_state1.get()
			check2 = chk_state2.get()
			if check1 == 1 or check2 == 1:
				program = main(file_name, saved_file_name)
				if check1 == 1:
					program[0]()
					messagebox.showinfo('Уведомление', 'Этап базового анализа выполнен')
					if check2 != 1:
						messagebox.showinfo('Уведомление', 'Работа завершена \nМожно выходить')
						if res == 'yes':
							show_results(saved_file_name)
				if check2 == 1:
					ngramm_window_show()
			elif check1 != 1 and not check2 != 1:
				messagebox.showinfo('Уведомление', 'Вы не выбрали ни одного параметра')
		except TypeError:
			messagebox.showinfo('Ошибка', 'Не указано имя файла для анализа')
	
	def load_results():
		messagebox.showinfo('Уведомление', 'Выберите папку для загрузки результатов')
		loading_directory = filedialog.askdirectory()
		show_results(loading_directory)
		
	def show_guide():
		guide_text_source = open('D:\\PITON\\guide_text.txt', "r", encoding='UTF-8')
		guide_text = guide_text_source.read()
		guide_window = Tk()
		guide_window.title('Инструкция по применению')
		guide_window_label = Label(guide_window, text=guide_text, justify='left')
		guide_window_label.grid(column=0, row=0)
	
	def get_file_name():
		messagebox.showinfo('Уведомление', 'Выберите файл для анализа')
		file_name = filedialog.askopenfilename()
		file_name_result = "{}".format(file_name)
		file_name_entry.delete(0, END)
		file_name_entry.insert(0, file_name_result)
		return file_name
	
	def get_save_directory():
		messagebox.showinfo('Уведомление', 'Выберите или создайте папку для сохранения результатов')
		saved_file_name = filedialog.askdirectory()
		res = "{}".format(saved_file_name)
		save_directory_entry.delete(0, END)
		save_directory_entry.insert(0, res)
		return saved_file_name
	
	def show_results(saved_file_name):
		result_window = Tk()
		result_window.title("Результаты анализа")
		tab_control = ttk.Notebook(result_window)
		tab1 = ttk.Frame(tab_control)
		tab2 = ttk.Frame(tab_control)
		tab3 = ttk.Frame(tab_control)
		tab_control.add(tab1, text='Результаты базового анализа')
		tab_control.add(tab2, text='Результаты анализа n-грамм')
		tab_control.add(tab3, text='Графики')
		tab_control.pack(expand=1, fill='both')
		base_analysis_scroll = scrolledtext.ScrolledText(tab1)
		ngram_analysis_scroll = scrolledtext.ScrolledText(tab2)
		try:
			base_analysis_file = docx.Document(f'{saved_file_name}\\Результаты анализа.docx')
			base_analysis_text = [paragraph.text for paragraph in base_analysis_file.paragraphs]
			base_analysis_text = "\n".join(base_analysis_text).partition('Г')[0]
			base_analysis_scroll.insert(INSERT, base_analysis_text)
		except FileNotFoundError:
			base_analysis_scroll.insert(INSERT, 'Анализ не выполнялся, либо файл был удален.\nПоказывать нечего')
		pd.set_option('display.max_rows', None)
		pd.set_option('display.max_columns', None)
		pd.set_option('display.max_colwidth', None)
		try:
			ngram_analysis_text = pd.read_excel(f'{saved_file_name}\\Результат подсчета n-грамм.xlsx')
			ngram_analysis_text = pd.DataFrame(ngram_analysis_text).astype(str)
			ngram_analysis_text = ngram_analysis_text.astype(str)
			ngram_analysis_scroll.insert(INSERT, ngram_analysis_text)
		except FileNotFoundError:
			ngram_analysis_scroll.insert(INSERT, 'Анализ n-грамм не выполнялся, либо файл был удален.\nПоказывать нечего')
		base_analysis_scroll.pack(expand=1, fill=Y)
		ngram_analysis_scroll.pack(expand=1, fill=Y)
		fig1 = plt.figure(figsize=(6.4, 4.8))
		plt.imshow(plt.imread(f'{saved_file_name}\\График 1. Слова до уборки стоп-слов.png'))
		plt.axis('off')
		canvas1 = FigureCanvasTkAgg(fig1, master=tab3)
		canvas1.get_tk_widget().pack(side=TOP)
		fig2 = plt.figure(figsize=(6.4, 4.8))
		plt.imshow(plt.imread(f'{saved_file_name}\\График 2. Слова после уборки стоп-слов.png'))
		plt.axis('off')
		canvas2 = FigureCanvasTkAgg(fig2, master=tab3)
		canvas2.get_tk_widget().pack(side=BOTTOM)
		result_window.mainloop()
		
	global window
	window = Tk()
	f_top = Frame(window)
	f_bot = Frame(window)
	f_left = Frame(window)
	f_right = Frame(window)
	window.title("Анализатор ЕЯ")
	window.resizable = True
	label = Label(f_top, text="Анализатор естественного языка", font=("Arial Bold", 14), justify='right')
	guide_button = Button(f_top, text='Инструкция по применению', command=show_guide, padx=5, pady=5, bd=0, fg='#fff',
	                      bg='#08f', underline=0, activebackground='#fff', activeforeground='#fff', cursor='hand2')
	load_button = Button(f_top, text="Загрузить прошлые результаты", command=load_results, padx=5, pady=5, bd=0, fg='#fff',
	                      bg='#08f', underline=0, activebackground='#fff', activeforeground='#fff', cursor='hand2')
	description1 = Label(f_left, text="Выберите составляющие анализа:", font=("Arial Bold", 14))
	description2 = Label(f_right, text="Выберите или введите путь к файлу", font=("Arial Bold", 14))
	description3 = Label(f_right, text="Выберите или введите путь\nдля сохранения результатов", font=("Arial Bold", 14))
	start_button = Button(f_left, text="ПУСК", command=start, padx=5, pady=5, bd=0, fg='#fff', bg='#08f',
	                      underline=0, activebackground='#fff', activeforeground='#fff', cursor='hand2')
	file_name_entry = Entry(f_right, width=50)
	file_name_entry.insert(INSERT, 'D:\\1.txt')
	save_directory_entry = Entry(f_right, width=50)
	save_directory_entry.insert(INSERT, 'D:\\результаты')
	file_name_button = Button(f_right, text="Загрузить", command=get_file_name, padx=5, pady=5, bd=0, fg='#fff',
	                          bg='#08f',
	                          underline=0, activebackground='#fff', activeforeground='#fff', cursor='hand2')
	save_directory_button = Button(f_right, text="Загрузить", command=get_save_directory, padx=5, pady=5, bd=0,
	                               fg='#fff', bg='#08f',
	                               underline=0, activebackground='#fff', activeforeground='#fff', cursor='hand2')
	chk_state1 = IntVar()
	chk_state1.set(1)
	chk1 = Checkbutton(f_left, text='Базовый анализ', var=chk_state1, cursor='hand2')
	chk_state2 = IntVar()
	chk_state2.set(1)
	chk2 = Checkbutton(f_left, text='Анализ n-грамм', var=chk_state2, cursor='hand2')
	
	f_top.pack(side=TOP)
	f_left.pack(side=LEFT)
	f_right.pack(side=RIGHT)
	f_bot.pack(side=BOTTOM, pady=10, padx=10)
	label.pack(side=TOP, padx=5, pady=5)
	load_button.pack(side=LEFT, padx=5, pady=5)
	guide_button.pack(side=TOP, padx=5, pady=15)
	description1.grid(column=0, row=0, padx=5, pady=5)
	chk1.grid(column=0, row=1, padx=5, pady=5)
	chk2.grid(column=0, row=2, padx=5, pady=5)
	description2.grid(column=0, row=0, padx=5, pady=5)
	file_name_entry.grid(column=0, row=1, padx=5, pady=5)
	file_name_button.grid(column=1, row=1, padx=5, pady=5)
	description3.grid(column=0, row=2, padx=5, pady=5)
	save_directory_entry.grid(column=0, row=3, padx=5, pady=5)
	save_directory_button.grid(column=1, row=3, padx=10, pady=10)
	start_button.grid(column=0, row=3, pady=20)
	
	def load_focus_in(e=None):
		load_button.configure(fg='#08f')
		load_button.configure(bg='#fff')
	
	def guide_focus_in(e=None):
		guide_button.configure(fg='#08f')
		guide_button.configure(bg='#fff')
	
	def start_focus_in(e=None):
		start_button.configure(fg='#08f')
		start_button.configure(bg='#fff')
	
	def file_focus_in(e=None):
		file_name_button.configure(fg='#08f')
		file_name_button.configure(bg='#fff')
	
	def save_focus_in(e=None):
		save_directory_button.configure(fg='#08f')
		save_directory_button.configure(bg='#fff')
	
	def load_focus_out(e=None):
		load_button.configure(bg='#08f')
		load_button.configure(fg='#fff')
	
	def guide_focus_out(e=None):
		guide_button.configure(bg='#08f')
		guide_button.configure(fg='#fff')
	
	def start_focus_out(e=None):
		start_button.configure(bg='#08f')
		start_button.configure(fg='#fff')
	
	def file_focus_out(e=None):
		file_name_button.configure(bg='#08f')
		file_name_button.configure(fg='#fff')
	
	def save_focus_out(e=None):
		save_directory_button.configure(bg='#08f')
		save_directory_button.configure(fg='#fff')
	
	load_button.bind('<Enter>', load_focus_in)
	load_button.bind('<Leave>', load_focus_out)
	guide_button.bind('<Enter>', guide_focus_in)
	guide_button.bind('<Leave>', guide_focus_out)
	start_button.bind('<Enter>', start_focus_in)
	start_button.bind('<Leave>', start_focus_out)
	file_name_button.bind('<Enter>', file_focus_in)
	file_name_button.bind('<Leave>', file_focus_out)
	save_directory_button.bind('<Enter>', save_focus_in)
	save_directory_button.bind('<Leave>', save_focus_out)
	
	window.mainloop()


interface()
